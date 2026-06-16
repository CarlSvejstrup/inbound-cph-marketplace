#!/usr/bin/env python3
"""Build the opstart-analyse DELIVERABLE: a .docx checklist report for a new client's
Google Ads account review (the 35-point Analysearbejdet).

WHY A .docx (not the audit slide deck): onboarding output lives in the client's Drive
folder as an editable document the specialist walks through and annotates. Slides are for
the client-facing audit (google-ads-general/ads-audit-report); this is the internal
opstarts-tjekliste. Decided with Carl 2026-06-10.

WHY A BUNDLED BUILDER (not the generic docx skill): Cowork plugins are self-contained.
This script self-bootstraps python-docx and runs on any machine with Python 3 + pip, same
pattern as optimering-loop/lib/review_workbook.py (openpyxl).

INPUT: a findings JSON (--in) with this shape (see SKILL.md Trin 4 for how it's filled):
{
  "client": "Dantaxi",
  "customer_id": "4149791707",
  "window": "Struktur-gennemgang (ny konto)" | "LAST_90_DAYS" | "...",
  "generated": "2026-06-10",            # passed in; this script never calls Date.now
  # NOTE: the overblik band (ok/warn/critical/no_data counts) is COMPUTED from the module
  # items below, never passed in. Any "summary" key in the input is ignored, so the band
  # can never disagree with the rendered rows.
  "headline_findings": ["...", "...", "..."],   # 3-5 top takeaways, Danish
  "modules": [
    {
      "key": "A", "title": "Annonceudvidelser (extensions)",
      "items": [
        {"n": 1, "punkt": "Sitelinks: min. 4 på hver kampagne",
         "status": "ok|warn|critical|no_data",
         "finding": "3 af 7 kampagner har <4 sitelinks: Brand, Generisk-DK, Lufthavn",
         # OPTIONAL depth fields — fill ONLY where they add value (concise points stay 1 line):
         "details": "Longer prose for modules that need depth (e.g. per-ad-group breakdown).",
         "evidence": [                         # verbatim location list — rendered as a "Hvor:" block
           "Kampagne 'IC | GSN | Hele DK' › ad group 'Aalborg' › RSA #447... › headline 'Bestil taxi til Aaborg' (skal være Aalborg)"
         ],
         "pointer": "for fuld RSA-gennemgang: kør optimering-loop (kræver at kontoen har kørt et stykke tid)."}
      ]
    }
  ],
  "sources": ["get_ad_extensions", "run_custom_gaql (campaign_asset)", "..."]
}

OUTPUT: a .docx at --out. Inbound house look: dark-navy headings, status chips as coloured
cells, one table per module, a summary band up top.

USAGE:
  python3 build_docx.py --in findings.json --out "Opstartsanalyse - Dantaxi - 2026-06-10.docx"
"""
from __future__ import annotations

import argparse
import json
import subprocess
import sys
from pathlib import Path


def _ensure_docx():
    """Import python-docx, installing it if absent. Tries a plain install first (the
    clean-venv case, e.g. Cowork), then falls back to --user and finally
    --break-system-packages so it also works on a PEP-668 externally-managed Python."""
    try:
        import docx  # noqa: F401
        return
    except ImportError:
        pass
    attempts = [
        ["--quiet", "python-docx"],
        ["--quiet", "--user", "python-docx"],
        ["--quiet", "--break-system-packages", "python-docx"],
    ]
    for flags in attempts:
        try:
            subprocess.run(
                [sys.executable, "-m", "pip", "install", *flags],
                check=True,
            )
            import docx  # noqa: F401
            return
        except (subprocess.CalledProcessError, ImportError):
            continue
    raise SystemExit(
        "Kunne ikke installere python-docx. Installer den manuelt: "
        "pip install python-docx (eller i et venv)."
    )


_ensure_docx()

from docx import Document  # noqa: E402
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL  # noqa: E402
from docx.enum.text import WD_ALIGN_PARAGRAPH  # noqa: E402
from docx.oxml.ns import qn  # noqa: E402
from docx.oxml import OxmlElement  # noqa: E402
from docx.shared import Pt, RGBColor, Inches  # noqa: E402

# --- Inbound design system (same palette as the audit deck / workbook) ---
NAVY = RGBColor(0x1A, 0x23, 0x40)        # dark navy headings/header band
NAVY_HEX = "1A2340"
WHITE = RGBColor(0xFF, 0xFF, 0xFF)
INK = RGBColor(0x22, 0x28, 0x33)         # body text
MUTED = RGBColor(0x6B, 0x72, 0x80)       # secondary text

# Status -> (label, cell fill hex, text colour)
STATUS = {
    "ok":       ("OK",            "DDF1E4", RGBColor(0x1E, 0x6B, 0x3F)),   # green tint
    "warn":     ("Kan forbedres", "FBF1D2", RGBColor(0x8A, 0x6D, 0x10)),   # amber tint
    "critical": ("Kritisk",       "F7DCDC", RGBColor(0x9B, 0x2C, 0x2C)),   # red tint
    "no_data":  ("Mangler data",  "ECEEF1", MUTED),                        # grey tint
}
STATUS_ORDER = ["critical", "warn", "ok", "no_data"]


def _count_statuses(modules: list) -> dict:
    """Tally the four status buckets across every item in every module. The overblik band
    is derived from this, so it can never contradict the rendered rows. Unknown statuses
    fall into no_data rather than vanishing."""
    counts = {"ok": 0, "warn": 0, "critical": 0, "no_data": 0}
    for mod in modules:
        for it in mod.get("items", []):
            status = it.get("status", "no_data")
            counts[status if status in counts else "no_data"] += 1
    return counts


def _is_handled(item: dict) -> bool:
    """Did the agent actually reach a verdict on this point? A point counts as HANDLED when
    it has any real status (ok/warn/critical) — i.e. the agent could assess it. no_data means
    it could not (missing history, no access), so it shows unchecked in the checklist. This is
    'did the agent walk the point', independent of whether the verdict was good or bad."""
    return item.get("status") in ("ok", "warn", "critical")


def _shade_cell(cell, hex_fill):
    """Set a table cell's background fill (python-docx has no direct API)."""
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_fill)
    tcPr.append(shd)


def _tight_row(row, height_pt=15):
    """Force a compact row: fixed (exact) height + vertically centered cells, and trim the
    cell paragraph spacing. Without this, Word auto-sizes rows tall and the checklist sprawls."""
    trPr = row._tr.get_or_add_trPr()
    th = OxmlElement("w:trHeight")
    th.set(qn("w:val"), str(int(height_pt * 20)))  # twips
    th.set(qn("w:hRule"), "atLeast")
    trPr.append(th)
    for cell in row.cells:
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        for p in cell.paragraphs:
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.space_after = Pt(0)
            p.paragraph_format.line_spacing = 1.0


def _set_cell_text(cell, text, *, color=INK, bold=False, size=10, align=None):
    cell.text = ""
    p = cell.paragraphs[0]
    if align is not None:
        p.alignment = align
    run = p.add_run(text)
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.color.rgb = color
    run.font.name = "Calibri"


def _heading(doc, text, *, size=15, color=NAVY, space_before=14, space_after=6):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after = Pt(space_after)
    run = p.add_run(text)
    run.font.size = Pt(size)
    run.font.bold = True
    run.font.color.rgb = color
    run.font.name = "Calibri"
    return p


def _set_col_widths(table, widths_inches):
    """Lock the table to fixed layout and write an explicit grid. Per-cell widths alone are
    ignored by Word when the table style autofits, so we set tblLayout=fixed AND rewrite the
    w:gridCol entries — that's what Word actually honours. Also set total tblW so it doesn't
    stretch to the margin."""
    twips = [int(w * 1440) for w in widths_inches]
    tblPr = table._tbl.tblPr
    # fixed layout
    for existing in tblPr.findall(qn("w:tblLayout")):
        tblPr.remove(existing)
    layout = OxmlElement("w:tblLayout")
    layout.set(qn("w:type"), "fixed")
    tblPr.append(layout)
    # total width
    for existing in tblPr.findall(qn("w:tblW")):
        tblPr.remove(existing)
    tblW = OxmlElement("w:tblW")
    tblW.set(qn("w:w"), str(sum(twips)))
    tblW.set(qn("w:type"), "dxa")
    tblPr.append(tblW)
    # rewrite the grid
    grid = table._tbl.find(qn("w:tblGrid"))
    if grid is not None:
        table._tbl.remove(grid)
    grid = OxmlElement("w:tblGrid")
    for tw in twips:
        gc = OxmlElement("w:gridCol")
        gc.set(qn("w:w"), str(tw))
        grid.append(gc)
    table._tbl.insert(list(table._tbl).index(tblPr) + 1, grid)
    # and per-cell width to match (belt and suspenders)
    table.autofit = False
    for row in table.rows:
        for cell, tw in zip(row.cells, twips):
            tcPr = cell._tc.get_or_add_tcPr()
            for existing in tcPr.findall(qn("w:tcW")):
                tcPr.remove(existing)
            tcW = OxmlElement("w:tcW")
            tcW.set(qn("w:w"), str(tw))
            tcW.set(qn("w:type"), "dxa")
            tcPr.append(tcW)


def build(findings: dict, out_path: str) -> str:
    doc = Document()

    # Base style
    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(10)
    normal.font.color.rgb = INK

    client = findings.get("client", "Ukendt klient")
    cid = findings.get("customer_id", "")
    window = findings.get("window", "")
    generated = findings.get("generated", "")

    # --- Title ---
    title = doc.add_paragraph()
    title.paragraph_format.space_after = Pt(2)
    r = title.add_run("Opstartsanalyse — Google Ads (Search)")
    r.font.size = Pt(22)
    r.font.bold = True
    r.font.color.rgb = NAVY
    r.font.name = "Calibri"

    sub = doc.add_paragraph()
    sub.paragraph_format.space_after = Pt(12)
    meta_bits = [client]
    if cid:
        meta_bits.append(f"konto {cid}")
    if generated:
        meta_bits.append(generated)
    rs = sub.add_run("  ·  ".join(meta_bits))
    rs.font.size = Pt(11)
    rs.font.color.rgb = MUTED
    rs.font.name = "Calibri"
    if window:
        wnd = doc.add_paragraph()
        wnd.paragraph_format.space_after = Pt(10)
        rw = wnd.add_run(f"Analysegrundlag: {window}")
        rw.font.size = Pt(9.5)
        rw.font.italic = True
        rw.font.color.rgb = MUTED

    # --- Summary band (4 status counts as a 1-row table) ---
    # Counts are COMPUTED from the module items, never read from the input. This is the
    # single source of truth: the band can never disagree with the rows below it. (A
    # caller-supplied 'summary' is intentionally ignored.)
    counts = _count_statuses(findings.get("modules", []))
    _heading(doc, "Overblik", size=14, space_before=4)
    sumt = doc.add_table(rows=2, cols=4)
    sumt.alignment = WD_TABLE_ALIGNMENT.LEFT
    labels = [
        ("critical", "Kritisk", counts["critical"]),
        ("warn", "Kan forbedres", counts["warn"]),
        ("ok", "OK", counts["ok"]),
        ("no_data", "Mangler data", counts["no_data"]),
    ]
    for i, (key, lbl, count) in enumerate(labels):
        _, fill, txt = STATUS[key]
        top = sumt.cell(0, i)
        _shade_cell(top, fill)
        _set_cell_text(top, str(count), color=txt, bold=True, size=20,
                       align=WD_ALIGN_PARAGRAPH.CENTER)
        bot = sumt.cell(1, i)
        _shade_cell(bot, fill)
        _set_cell_text(bot, lbl, color=txt, bold=False, size=9,
                       align=WD_ALIGN_PARAGRAPH.CENTER)
    _set_col_widths(sumt, [1.6, 1.6, 1.6, 1.6])

    # --- Headline findings ---
    headlines = findings.get("headline_findings", [])
    if headlines:
        _heading(doc, "Vigtigste fund", size=14)
        for h in headlines:
            b = doc.add_paragraph(style="List Bullet")
            rb = b.add_run(h)
            rb.font.size = Pt(10.5)
            rb.font.color.rgb = INK

    # --- Checklist / indholdsfortegnelse (coverage proof) ---
    # The full point list with a ✓ / ☐ box per point, showing what the agent actually
    # handled (reached a verdict on) vs. couldn't assess. This mirrors the ClickUp subtask
    # list inside the doc, so the specialist sees coverage at a glance. The VERDICT (OK/
    # warn/critical) is NOT shown here — that lives in the detail tables below. This box is
    # purely "did the agent walk the point". TWO checkboxes per point: an Agent column
    # (filled by the skill — what the agent handled) and an Ekspert column (always an empty ☐
    # for the specialist to tick by hand after reviewing). The agent proposes, the expert
    # disposes — same human-in-the-loop pattern as the ClickUp subtask getting its real
    # checkmark from a human, but inside the document. Layout per module: Agent | Ekspert | punkt.
    modules = findings.get("modules", [])
    all_items = [it for mod in modules for it in mod.get("items", [])]
    handled = sum(1 for it in all_items if _is_handled(it))
    total = len(all_items)
    _heading(doc, "Tjekliste — hvad er gennemgået", size=14, space_before=16)
    cov = doc.add_paragraph()
    cov.paragraph_format.space_after = Pt(6)
    rc = cov.add_run(
        f"Agenten har behandlet {handled} af {total} punkter. "
        "Kolonnen Agent (✓ = gennemgået · ☐ = kunne ikke vurderes) er agentens egen registrering. "
        "Kolonnen Ekspert er tom til at sætte flueben i hånden, når du har gennemgået fundet. "
        "Selve vurderingen står i modul-sektionerne nedenfor."
    )
    rc.font.size = Pt(9)
    rc.font.italic = True
    rc.font.color.rgb = MUTED
    # ONE compact table for all points (not a table per module — that sprawled). Columns:
    # Modul | # | Tjekpunkt | Agent | Ekspert. The module name prints only on its first row,
    # so the grouping reads without repeating headers or leaving big gaps.
    ct = doc.add_table(rows=1, cols=5)
    ct.alignment = WD_TABLE_ALIGNMENT.LEFT
    ct.style = "Table Grid"
    hdr = ct.rows[0]
    for j, htext in enumerate(["Modul", "#", "Tjekpunkt", "Agent", "Ekspert"]):
        hc = hdr.cells[j]
        _shade_cell(hc, NAVY_HEX)
        _set_cell_text(hc, htext, color=WHITE, bold=True, size=8.5,
                       align=(WD_ALIGN_PARAGRAPH.CENTER if htext in ("Agent", "Ekspert", "#")
                              else WD_ALIGN_PARAGRAPH.LEFT))
    _tight_row(hdr, height_pt=16)
    for mod in modules:
        items = mod.get("items", [])
        if not items:
            continue
        mod_label = mod.get("key", "")  # short — "A", "B", … keeps the Modul column narrow
        for idx, it in enumerate(items):
            done = _is_handled(it)
            row = ct.add_row()
            # module key only on the group's first row
            _set_cell_text(row.cells[0], mod_label if idx == 0 else "",
                           color=NAVY, bold=True, size=9, align=WD_ALIGN_PARAGRAPH.CENTER)
            # number
            _set_cell_text(row.cells[1], str(it.get("n", "")),
                           color=NAVY, bold=True, size=9, align=WD_ALIGN_PARAGRAPH.CENTER)
            # point text
            _set_cell_text(row.cells[2], it.get("punkt", ""),
                           color=(INK if done else MUTED), bold=False, size=9)
            # agent box (filled)
            _set_cell_text(row.cells[3], "✓" if done else "☐",
                           color=(STATUS["ok"][2] if done else MUTED),
                           bold=True, size=10.5, align=WD_ALIGN_PARAGRAPH.CENTER)
            # expert box (always empty)
            _set_cell_text(row.cells[4], "☐", color=MUTED, bold=False, size=10.5,
                           align=WD_ALIGN_PARAGRAPH.CENTER)
            _tight_row(row, height_pt=14)
    _set_col_widths(ct, [0.55, 0.35, 4.75, 0.58, 0.62])

    # --- Per-module detail tables ---
    _heading(doc, "Detaljer per modul", size=14, space_before=18)
    det = doc.add_paragraph()
    det.paragraph_format.space_after = Pt(4)
    rd = det.add_run("Status og konkret fund for hvert punkt.")
    rd.font.size = Pt(9)
    rd.font.italic = True
    rd.font.color.rgb = MUTED
    for mod in findings.get("modules", []):
        mtitle = mod.get("title", mod.get("key", ""))
        _heading(doc, mtitle, size=12, space_before=14)
        items = mod.get("items", [])
        if not items:
            continue
        # 3 columns: Status | # + Punkt | Fund
        tbl = doc.add_table(rows=1, cols=3)
        tbl.alignment = WD_TABLE_ALIGNMENT.LEFT
        tbl.style = "Table Grid"
        # header row
        hdr = tbl.rows[0]
        for j, htext in enumerate(["Status", "Tjekpunkt", "Fund"]):
            c = hdr.cells[j]
            _shade_cell(c, NAVY_HEX)
            _set_cell_text(c, htext, color=WHITE, bold=True, size=9.5)
        for it in items:
            status = it.get("status", "no_data")
            label, fill, txt = STATUS.get(status, STATUS["no_data"])
            row = tbl.add_row()
            # status cell
            sc = row.cells[0]
            _shade_cell(sc, fill)
            _set_cell_text(sc, label, color=txt, bold=True, size=9)
            # punkt cell (number + question)
            pc = row.cells[1]
            n = it.get("n", "")
            punkt = it.get("punkt", "")
            pc.text = ""
            pp = pc.paragraphs[0]
            rn = pp.add_run(f"{n}. " if n else "")
            rn.font.size = Pt(9.5)
            rn.font.bold = True
            rn.font.color.rgb = NAVY
            rt = pp.add_run(punkt)
            rt.font.size = Pt(9.5)
            rt.font.color.rgb = INK
            # finding cell: prose, then optional longer `details`, then a verbatim
            # `evidence` "Hvor:" location block, then an optional `pointer` line.
            fc = row.cells[2]
            fc.text = ""
            fp = fc.paragraphs[0]
            fp.paragraph_format.space_before = Pt(0)
            fp.paragraph_format.space_after = Pt(0)
            rf = fp.add_run(it.get("finding", ""))
            rf.font.size = Pt(9.5)
            rf.font.color.rgb = INK

            # details: optional longer prose (only modules that need depth fill this)
            details = it.get("details", "")
            if details:
                dp = fc.add_paragraph()
                dp.paragraph_format.space_before = Pt(3)
                dp.paragraph_format.space_after = Pt(0)
                rdd = dp.add_run(details)
                rdd.font.size = Pt(9)
                rdd.font.color.rgb = INK

            # evidence: a verbatim location list — campaign / ad group / ad / the offending
            # string. Rendered as "Hvor:" + one line per item so the specialist can act.
            evidence = it.get("evidence", [])
            if evidence:
                ep = fc.add_paragraph()
                ep.paragraph_format.space_before = Pt(3)
                ep.paragraph_format.space_after = Pt(0)
                rlbl = ep.add_run("Hvor: ")
                rlbl.font.size = Pt(9)
                rlbl.font.bold = True
                rlbl.font.color.rgb = NAVY
                for ei, ev in enumerate(evidence):
                    line = fc.add_paragraph() if ei else ep
                    if ei:
                        line.paragraph_format.space_before = Pt(0)
                        line.paragraph_format.space_after = Pt(0)
                    rev = line.add_run(("• " if ei else "") + str(ev))
                    rev.font.size = Pt(9)
                    rev.font.color.rgb = MUTED

            # pointer: optional "for the full deep-dive, run X" line
            pointer = it.get("pointer", "")
            if pointer:
                pp2 = fc.add_paragraph()
                pp2.paragraph_format.space_before = Pt(3)
                pp2.paragraph_format.space_after = Pt(0)
                rp = pp2.add_run("→ " + pointer)
                rp.font.size = Pt(9)
                rp.font.italic = True
                rp.font.color.rgb = NAVY
            _tight_row(row, height_pt=12)
        _set_col_widths(tbl, [0.95, 2.05, 3.7])

    # --- Sources footer ---
    sources = findings.get("sources", [])
    if sources:
        _heading(doc, "Datakilder", size=12, space_before=18)
        src = doc.add_paragraph()
        rsrc = src.add_run(", ".join(sources))
        rsrc.font.size = Pt(9)
        rsrc.font.italic = True
        rsrc.font.color.rgb = MUTED

    note = doc.add_paragraph()
    note.paragraph_format.space_before = Pt(10)
    rnote = note.add_run(
        "Read-only mod kontoen. Ingen ændringer er foretaget. "
        "Status er en struktur- og hygiejne-vurdering, ikke en performance-dom; "
        "punkter uden historik er markeret “Mangler data”."
    )
    rnote.font.size = Pt(8.5)
    rnote.font.italic = True
    rnote.font.color.rgb = MUTED

    out = Path(out_path)
    doc.save(str(out))
    return str(out)


def _validate(findings: dict) -> list[str]:
    """Cheap structural checks so a malformed findings object fails loudly, not silently."""
    errs = []
    if not findings.get("client"):
        errs.append("mangler 'client'")
    mods = findings.get("modules")
    if not isinstance(mods, list) or not mods:
        errs.append("mangler 'modules' (skal være en ikke-tom liste)")
    else:
        seen = 0
        for m in mods:
            for it in m.get("items", []):
                seen += 1
                if it.get("status") not in STATUS:
                    errs.append(
                        f"punkt {it.get('n')}: ugyldig status {it.get('status')!r} "
                        f"(skal være en af {list(STATUS)})"
                    )
        if seen == 0:
            errs.append("ingen tjekpunkter fundet i 'modules'")
    return errs


def main():
    ap = argparse.ArgumentParser(description="Build opstart-analyse .docx report")
    ap.add_argument("--in", dest="infile", required=True, help="findings JSON path")
    ap.add_argument("--out", dest="outfile", required=True, help="output .docx path")
    args = ap.parse_args()

    findings = json.loads(Path(args.infile).read_text(encoding="utf-8"))
    errs = _validate(findings)
    if errs:
        print("FINDINGS-VALIDERING FEJLEDE:", file=sys.stderr)
        for e in errs:
            print(f"  - {e}", file=sys.stderr)
        sys.exit(2)

    out = build(findings, args.outfile)
    print(out)


if __name__ == "__main__":
    main()
